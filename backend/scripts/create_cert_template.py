"""
One-time script to generate backend/templates/comp_cert_template.docx.
Run from the project root: python backend/scripts/create_cert_template.py

The template uses docx-templates syntax with cmdDelimiter ['{{', '}}']:
  - Simple values:  {{field_name}}
  - Loop start:     {{FOR row IN rows}}  (in first cell of template row)
  - Loop end:       {{END-FOR row}}      (in last cell of same row)
  - Loop field ref: {{row.field}}
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(os.path.dirname(__file__), '..', 'templates', 'comp_cert_template.docx')


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_field(doc, label, placeholder):
    """Add a labelled field line: 'Label: {{placeholder}}'"""
    p = doc.add_paragraph()
    run_label = p.add_run(label + ': ')
    run_label.bold = True
    p.add_run('{{' + placeholder + '}}')
    return p


def add_page_break(doc):
    doc.add_page_break()


def make_table(doc, headers, loop_var, loop_array, field_names):
    """
    Add a table with docx-templates FOR loop structure:
      Row 0: column headers
      Row 1: {{FOR loop_var IN loop_array}}   ← marker row, removed from output
      Row 2: {{loop_var.field1}} | ... | {{loop_var.fieldN}}  ← repeated per item
      Row 3: {{END-FOR loop_var}}             ← marker row, removed from output
    """
    n_cols = len(field_names)
    tbl = doc.add_table(rows=4, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Row 0: headers
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    # Row 1: FOR marker (first cell only, rest blank)
    tbl.rows[1].cells[0].text = '{{FOR ' + loop_var + ' IN ' + loop_array + '}}'

    # Row 2: data fields — loop var must be prefixed with $ inside the loop
    for i, fname in enumerate(field_names):
        tbl.rows[2].cells[i].text = '{{$' + loop_var + '.' + fname + '}}'

    # Row 3: END-FOR marker (first cell only, rest blank)
    tbl.rows[3].cells[0].text = '{{END-FOR ' + loop_var + '}}'

    return tbl


doc = Document()

# ── Default styles ────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── PAGE 1: INTRO ─────────────────────────────────────────────────────────────
add_heading(doc, 'COMPENSATION CERTIFICATE', level=1)
add_heading(doc, 'Ngonye Falls Hydropower Project', level=2)

doc.add_paragraph()
add_field(doc, 'PAH Number', 'pah')
add_field(doc, 'Household Head', 'head_name')
add_field(doc, 'NRC', 'head_nrc')
add_field(doc, 'Village', 'village')
add_field(doc, 'Survey Date', 'survey_date')
add_field(doc, 'Certificate Date', 'cert_date')

add_page_break(doc)

# ── PAGE 2: LAND CHOICES ──────────────────────────────────────────────────────
add_heading(doc, 'Land Choices', level=1)

doc.add_paragraph()
add_field(doc, 'Landholding Option', 'icaoption_landholding')
add_field(doc, 'Dry Land Option', 'icaoption_dryland')
add_field(doc, 'Garden Option', 'icaoption_garden')
add_field(doc, 'Structure Location Option', 'icaoption_structure_location')

add_page_break(doc)

# ── PAGE 3: LAND SPEC ─────────────────────────────────────────────────────────
add_heading(doc, 'Land Specification', level=1)

doc.add_paragraph()
make_table(
    doc,
    headers=['Land Class', 'Acquisition', 'Area (m²)', 'Rate (ZMW/m²)', 'Prep Allowance', 'Land Value (ZMW)', 'Compensation Option', 'Remaining Viable', 'Irrigation Eligible'],
    loop_var='row',
    loop_array='land_rows',
    field_names=['land_class', 'acquisition_class', 'area_sqm', 'rate_acquisition', 'prep_allowance', 'land_value', 'compensation_option', 'remaining_viable', 'irrigation_eligible'],
)

add_page_break(doc)

# ── PAGE 4: PRIMARY STRUCTURES ────────────────────────────────────────────────
add_heading(doc, 'Primary Structures', level=1)

doc.add_paragraph()
make_table(
    doc,
    headers=['Structure Type', 'Walls', 'Roof', 'Floor', 'Replacement Option', 'Value (ZMW)'],
    loop_var='ps',
    loop_array='primary_structure_rows',
    field_names=['structure_type', 'walls_type', 'roof_type', 'floor_type', 'replacement_option', 'structure_value'],
)

add_page_break(doc)

# ── PAGE 5: SECONDARY STRUCTURES ─────────────────────────────────────────────
add_heading(doc, 'Secondary Structures', level=1)

doc.add_paragraph()
make_table(
    doc,
    headers=['Description', 'Structure Type', 'Walls', 'Roof', 'Floor', 'Value (ZMW)'],
    loop_var='ss',
    loop_array='secondary_structure_rows',
    field_names=['secondary_description', 'structure_type', 'walls_type', 'roof_type', 'floor_type', 'structure_value'],
)

add_page_break(doc)

# ── PAGE 6: CROPS ─────────────────────────────────────────────────────────────
add_heading(doc, 'Crops', level=1)

doc.add_paragraph()
make_table(
    doc,
    headers=['Crop Type', 'Area', 'Rate (ZMW/ha)', 'Value (ZMW)'],
    loop_var='cr',
    loop_array='crop_rows',
    field_names=['crop_type', 'crop_size', 'rate', 'crop_value'],
)
add_field(doc, 'Total Crop Value (ZMW)', 'total_crop_value')

add_page_break(doc)

# ── PAGE 7: TREES ─────────────────────────────────────────────────────────────
add_heading(doc, 'Trees', level=1)

doc.add_paragraph()
make_table(
    doc,
    headers=['Tree Type', 'Productive', 'Juvenile', 'Replacement Saplings', 'Compensation (ZMW)'],
    loop_var='tr',
    loop_array='tree_rows',
    field_names=['tree_type', 'productive_count', 'juvenile_count', 'replacement_saplings', 'compensation'],
)
doc.add_paragraph()
add_field(doc, 'Total Productive Trees', 'total_productive')
add_field(doc, 'Total Juvenile Trees', 'total_juvenile')
add_field(doc, 'Total Replacement Saplings', 'total_saplings')
add_field(doc, 'Total Tree Compensation (ZMW)', 'total_trees_comp')

add_page_break(doc)

# ── PAGE 8: LR / COMMUNAL ────────────────────────────────────────────────────
add_heading(doc, 'Livelihood Restoration / Communal', level=1)

doc.add_paragraph()
add_heading(doc, 'Livelihood Restoration Entitlements', level=2)
add_field(doc, 'Agricultural Development', 'lr_agricultural')
add_field(doc, 'Livestock Development', 'lr_livestock')
add_field(doc, 'Water and Sanitation Support', 'lr_water')
add_field(doc, 'Fisheries Support', 'lr_fisheries')
add_field(doc, 'Reed Beds Support', 'lr_reedbeds')
add_field(doc, 'Agriculture Inputs / Credit / Enterprise Development', 'lr_agricultureinputs')

doc.add_paragraph()
add_heading(doc, 'Graves', level=2)
add_field(doc, 'Grave Visit Option', 'graves_visit')
add_field(doc, 'Grave Relocation Option', 'graves_option')
add_field(doc, 'Ceremony Arrangement', 'graves_ceremony')

add_page_break(doc)

# ── PAGE 9: ALLOWANCES ────────────────────────────────────────────────────────
add_heading(doc, 'Allowances and Compensation Summary', level=1)

doc.add_paragraph()
add_heading(doc, 'Cash Allowances', level=2)
add_field(doc, 'Disturbance Allowance (ZMW)', 'allowance_disturbance')
add_field(doc, 'Rental Allowance (ZMW)', 'allowance_rental')
add_field(doc, 'Transport Allowance (ZMW)', 'allowance_transport')
add_field(doc, 'Transitional Allowance (ZMW)', 'allowance_transitional')
add_field(doc, 'Business Allowance (ZMW)', 'allowance_business')
add_field(doc, 'Land Preparation Allowance (ZMW)', 'allowance_landprep')
add_field(doc, 'Total Allowances (ZMW)', 'allowance_total')

doc.add_paragraph()
add_heading(doc, 'Compensation Components', level=2)
add_field(doc, 'Structures Compensation (ZMW)', 'structures_compensation')
add_field(doc, 'Land Compensation (ZMW)', 'land_compensation_value')
add_field(doc, 'Tree Compensation (ZMW)', 'total_trees_comp')
add_field(doc, 'Crop Compensation (ZMW)', 'total_crop_value')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('GRAND TOTAL (ZMW): ')
run.bold = True
run.font.size = Pt(12)
val_run = p.add_run('{{grand_total}}')
val_run.bold = True
val_run.font.size = Pt(12)

add_page_break(doc)

# ── PAGE 10: SIGNATURES ───────────────────────────────────────────────────────
add_heading(doc, 'Acknowledgement and Signatures', level=1)

doc.add_paragraph(
    'I/We, the undersigned, confirm that the entitlements described in this '
    'Compensation Certificate have been explained to me/us and that I/we agree '
    'to the compensation and livelihood restoration provisions set out herein.'
)

doc.add_paragraph()
add_heading(doc, 'Household Head', level=2)
add_field(doc, 'Name', 'sig_head_name')
add_field(doc, 'NRC', 'sig_head_nrc')
doc.add_paragraph('Signature: _______________________________')
doc.add_paragraph('Date: ___________________________________')

doc.add_paragraph()
add_heading(doc, 'Co-signatory', level=2)
add_field(doc, 'Name', 'sig_cosig_name')
add_field(doc, 'NRC', 'sig_cosig_nrc')
doc.add_paragraph('Signature: _______________________________')
doc.add_paragraph('Date: ___________________________________')

doc.add_paragraph()
add_heading(doc, 'Western Power Corporation Representative', level=2)
doc.add_paragraph('Name: ___________________________________')
doc.add_paragraph('Signature: _______________________________')
add_field(doc, 'ICA Date', 'ica_date')

# ── DEBUG: Data Object ────────────────────────────────────────────────────────
add_page_break(doc)
add_heading(doc, 'DEBUG — Data Object', level=1)
doc.add_paragraph('Generated automatically. Remove this section before issuing certificates.')

# One table row per JSON line: FOR marker | line | END-FOR marker
tbl = doc.add_table(rows=3, cols=1)
tbl.style = 'Table Grid'
tbl.rows[0].cells[0].text = '{{FOR line IN debug_lines}}'
tbl.rows[1].cells[0].text = '{{$line}}'
tbl.rows[2].cells[0].text = '{{END-FOR line}}'
# Set monospace font on the data row
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement as _OxmlEl
run = tbl.rows[1].cells[0].paragraphs[0].runs[0]
run.font.name = 'Courier New'
run.font.size = Pt(8)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.abspath(OUT_PATH)
doc.save(out)
print(f'Template written to: {out}')
